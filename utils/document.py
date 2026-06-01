import docx
import pandas as pd
import re

def extract_paragraphs(file_path):
    """Reads docx and returns a list of paragraph texts."""
    doc = docx.Document(file_path)
    return [p.text.strip() for p in doc.paragraphs]

def split_sentences(paragraphs, active_paragraphs, api_key=None, split_model=None):
    """
    Splits active paragraphs into sentences.
    active_paragraphs is a list of booleans corresponding to paragraphs.
    Returns: sentences list, sentence_mapping list (maps sentence index to paragraph index)
    """
    sentences = []
    sentence_mapping = [] 
    
    for p_idx, text in enumerate(paragraphs):
        if not text or not active_paragraphs[p_idx]:
            continue
            
        if split_model and split_model != "Mock Test Mode":
            from utils.llm import split_sentences_llm
            p_sentences = split_sentences_llm(api_key, split_model, text)
        else:
            # Fallback regex splitter
            p_sentences = [s.strip() for s in re.split(r'(?<=[.!?])\s+', text) if s.strip()]
            
        for s in p_sentences:
            sentences.append(s)
            sentence_mapping.append(p_idx)
            
    return sentences, sentence_mapping

def create_revised_docx(original_file_path, paragraphs, sentences, sentence_mapping, finalized_sentences, active_paragraphs, output_path):
    # Start with original paragraphs
    new_paragraphs = list(paragraphs)
    
    # We clear the active ones because they will be rebuilt from finalized_sentences
    for p_idx, active in enumerate(active_paragraphs):
        if active:
            new_paragraphs[p_idx] = ""
            
    # Rebuild active paragraphs from finalized sentences
    for s_idx, final_s in enumerate(finalized_sentences):
        p_idx = sentence_mapping[s_idx]
        if new_paragraphs[p_idx]:
            new_paragraphs[p_idx] += " " + final_s
        else:
            new_paragraphs[p_idx] = final_s
            
    # Load original to preserve overall document styles, but rewrite paragraph content
    doc = docx.Document(original_file_path)
    
    for p_idx, p in enumerate(doc.paragraphs):
        p.clear()
        if p_idx < len(new_paragraphs):
            p.add_run(new_paragraphs[p_idx])
            
    doc.save(output_path)
    
def create_integrity_report(history_data, output_path):
    df = pd.DataFrame(history_data)
    if 'Suggestions' in df.columns:
        df['Suggestions'] = df['Suggestions'].apply(lambda x: "\n".join(x) if isinstance(x, list) else x)
    df.to_excel(output_path, index=False)
